"""
DOCX Generator - Creates Word documents
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, Any, List
import os
from datetime import datetime

class DocxGenerator:
    """
    Generates professional Word documents (.docx)
    """
    
    def __init__(self):
        """Initialize the DOCX generator"""
        self.output_dir = "output"
        self._ensure_output_dir()
        print("✓ DocxGenerator initialized")
    
    def _ensure_output_dir(self):
        """Create output directory if it doesn't exist"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def create_document(self, content: Dict[str, Any]) -> str:
        """
        Create a Word document from content
        
        Args:
            content: {
                'title': 'Document title',
                'subtitle': 'Optional subtitle',
                'sections': [
                    {
                        'heading': 'Section heading',
                        'content': 'Section content text',
                        'type': 'paragraph' | 'bullet_list' | 'numbered_list'
                    }
                ],
                'metadata': {
                    'author': 'Author name',
                    'company': 'Company name',
                    'date': 'Date string'
                }
            }
        
        Returns:
            Path to created document
        """
        
        print(f"\n📄 Creating Word document: {content.get('title', 'Untitled')}")
        
        # Create document
        doc = Document()
        
        # Set up styles
        self._setup_styles(doc)
        
        # Add title page
        self._add_title_page(doc, content)
        
        # Add page break
        doc.add_page_break()
        
        # Add sections
        sections = content.get('sections', [])
        for i, section in enumerate(sections):
            self._add_section(doc, section)
            
            # Add space between sections (but not after last one)
            if i < len(sections) - 1:
                doc.add_paragraph()
        
        # Add footer
        self._add_footer(doc, content)
        
        # Save document
        filename = self._generate_filename(content.get('title', 'document'))
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        print(f"✓ Document created: {filepath}")
        
        return filepath
    
    def _setup_styles(self, doc: Document):
        """Set up document styles"""
        
        # Create custom heading style
        styles = doc.styles
        
        # Title style
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(28)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    def _add_title_page(self, doc: Document, content: Dict[str, Any]):
        """Add title page to document"""
        
        # Add title
        title = content.get('title', 'Untitled Document')
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.size = Pt(28)
        title_para.runs[0].font.bold = True
        title_para.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Add subtitle if exists
        subtitle = content.get('subtitle', '')
        if subtitle:
            doc.add_paragraph()  # Spacing
            subtitle_para = doc.add_paragraph(subtitle)
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_para.runs[0].font.size = Pt(16)
            subtitle_para.runs[0].font.color.rgb = RGBColor(89, 89, 89)
        
        # Add metadata
        metadata = content.get('metadata', {})
        if metadata:
            doc.add_paragraph()  # Spacing
            doc.add_paragraph()
            doc.add_paragraph()
            
            meta_lines = []
            if 'author' in metadata:
                meta_lines.append(f"Author: {metadata['author']}")
            if 'company' in metadata:
                meta_lines.append(f"Company: {metadata['company']}")
            if 'date' in metadata:
                meta_lines.append(f"Date: {metadata['date']}")
            else:
                meta_lines.append(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            
            for line in meta_lines:
                meta_para = doc.add_paragraph(line)
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                meta_para.runs[0].font.size = Pt(11)
    
    def _add_section(self, doc: Document, section: Dict[str, Any]):
        """Add a section to the document"""
        
        # Add heading
        heading = section.get('heading', '')
        if heading:
            heading_para = doc.add_heading(heading, level=1)
            heading_para.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Add content based on type
        content_text = section.get('content', '')
        content_type = section.get('type', 'paragraph')
        
        if content_type == 'paragraph':
            para = doc.add_paragraph(content_text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        elif content_type == 'bullet_list':
            # Split content into bullet points
            items = content_text.split('\n')
            for item in items:
                if item.strip():
                    doc.add_paragraph(item.strip(), style='List Bullet')
        
        elif content_type == 'numbered_list':
            # Split content into numbered items
            items = content_text.split('\n')
            for item in items:
                if item.strip():
                    doc.add_paragraph(item.strip(), style='List Number')
    
    def _add_footer(self, doc: Document, content: Dict[str, Any]):
        """Add footer with page numbers"""
        
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated by DocFlow AI | {datetime.now().strftime('%B %d, %Y')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(9)
        footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    def _generate_filename(self, title: str) -> str:
        """Generate a safe filename from title"""
        
        # Clean title for filename
        safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_') else '' for c in title)
        safe_title = safe_title.replace(' ', '_')
        
        # Add timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        return f"{safe_title}_{timestamp}.docx"
    
    def create_simple_document(self, title: str, content: str) -> str:
        """
        Quick method to create a simple document
        
        Args:
            title: Document title
            content: Main content text
            
        Returns:
            Path to created document
        """
        
        doc_content = {
            'title': title,
            'sections': [
                {
                    'heading': '',
                    'content': content,
                    'type': 'paragraph'
                }
            ],
            'metadata': {
                'author': 'DocFlow AI',
                'date': datetime.now().strftime('%B %d, %Y')
            }
        }
        
        return self.create_document(doc_content)