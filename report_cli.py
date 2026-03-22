#!/usr/bin/env python3
"""
Infopercept Report Generator CLI
---------------------------------
Usage:
    python report_cli.py "create a sales report"
    python report_cli.py "generate hr report"
    python report_cli.py "make a finance report"
    python report_cli.py   (interactive mode — prompts for input)
"""

import sys
import os
import json
import re
import textwrap
from datetime import datetime
from pathlib import Path

# ─── Colour helpers (works on any ANSI terminal) ────────────────────────────
RESET  = "\033[0m"
BOLD   = "\033[1m"
DIM    = "\033[2m"
RED    = "\033[91m"
GREEN  = "\033[92m"
YELLOW = "\033[93m"
CYAN   = "\033[96m"
WHITE  = "\033[97m"
BLUE   = "\033[94m"
MAGENTA= "\033[95m"

def c(text, *codes): return "".join(codes) + str(text) + RESET
def banner():
    print()
    print(c("╔══════════════════════════════════════════════════════════════╗", CYAN, BOLD))
    print(c("║", CYAN, BOLD) + c("     INFOPERCEPT  —  AI Report Generator CLI  v1.0          ", WHITE, BOLD) + c("║", CYAN, BOLD))
    print(c("║", CYAN, BOLD) + c("     Infopercept Consulting Pvt. Ltd.                        ", DIM)         + c("║", CYAN, BOLD))
    print(c("╚══════════════════════════════════════════════════════════════╝", CYAN, BOLD))
    print()

def step(icon, msg):  print(f"  {icon}  {msg}")
def ok(msg):          print(c(f"  ✔  {msg}", GREEN))
def warn(msg):        print(c(f"  ⚠  {msg}", YELLOW))
def err(msg):         print(c(f"  ✖  {msg}", RED))
def info(msg):        print(c(f"  ℹ  {msg}", CYAN))
def divider():        print(c("  " + "─"*60, DIM))
print("hi")
# ─── Data layer ─────────────────────────────────────────────────────────────
BASE_DIR  = Path(__file__).parent
DATA_DIR  = BASE_DIR / "data"
OUT_DIR   = BASE_DIR / "outputs"
OUT_DIR.mkdir(exist_ok=True)

DATA_CATALOGUE = {
    "sales": {
        "file":     DATA_DIR / "sales_2024.json",
        "keywords": ["sales", "revenue", "deals", "pipeline", "quota", "win rate",
                     "regional", "product", "customers", "upsell"],
        "label":    "Sales Performance",
        "emoji":    "📊",
    },
    "hr": {
        "file":     DATA_DIR / "hr_2024.json",
        "keywords": ["hr", "human resources", "people", "headcount", "hiring",
                     "attrition", "employee", "workforce", "talent", "diversity",
                     "compensation", "payroll", "engagement"],
        "label":    "Human Resources",
        "emoji":    "👥",
    },
    "finance": {
        "file":     DATA_DIR / "finance_2024.json",
        "keywords": ["finance", "financial", "revenue", "profit", "ebitda",
                     "balance sheet", "cash flow", "income", "p&l", "margin",
                     "earnings", "fiscal"],
        "label":    "Financial Performance",
        "emoji":    "💰",
    },
}

def detect_report_type(prompt: str) -> str | None:
    """Score each data source against the user prompt and return the best match."""
    prompt_lower = prompt.lower()
    scores = {}
    for key, meta in DATA_CATALOGUE.items():
        score = sum(1 for kw in meta["keywords"] if kw in prompt_lower)
        # Direct name match is a strong signal
        if key in prompt_lower or meta["label"].lower() in prompt_lower:
            score += 10
        scores[key] = score
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else None

def load_data(report_type: str) -> dict:
    meta = DATA_CATALOGUE[report_type]
    with open(meta["file"]) as f:
        return json.load(f)

# ─── LLM layer (Anthropic API via raw HTTP) ─────────────────────────────────
def call_llm(system: str, user: str, api_key: str) -> str:
    import urllib.request, urllib.error
    payload = json.dumps({
        "model": "claude-sonnet-4-20250514",
        "max_tokens": 2048,
        "system": system,
        "messages": [{"role": "user", "content": user}]
    }).encode()
    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
        },
        method="POST"
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            body = json.loads(resp.read())
            return body["content"][0]["text"]
    except urllib.error.HTTPError as e:
        raise RuntimeError(f"API error {e.code}: {e.read().decode()}")

def get_api_key() -> str | None:
    return os.environ.get("ANTHROPIC_API_KEY")

def generate_analysis(data: dict, prompt: str, api_key: str) -> str:
    """Ask the LLM to write a deep analytical narrative from the raw data."""
    system = textwrap.dedent("""
        You are a senior business analyst at Infopercept Consulting Pvt. Ltd.
        You receive structured JSON business data and write crisp, professional
        report narratives. Write in clear business English. Use specific numbers
        from the data. No bullet points — full paragraphs only.
        Your output will be inserted verbatim into a formal report.
    """).strip()

    user = textwrap.dedent(f"""
        User request: "{prompt}"

        Data:
        {json.dumps(data, indent=2)}

        Write the following sections (use these exact headings):
        ## Executive Summary
        ## Key Highlights
        ## Detailed Analysis
        ## Strategic Recommendations

        Each section should be 2-4 sentences and reference actual numbers.
        Keep total output under 600 words.
    """).strip()

    return call_llm(system, user, api_key)

# ─── Formatters ─────────────────────────────────────────────────────────────

def fmt_usd(n):
    if n >= 1_000_000: return f"${n/1_000_000:.2f}M"
    if n >= 1_000:     return f"${n/1_000:.0f}K"
    return f"${n:,}"

def today_str():
    return datetime.now().strftime("%B %d, %Y")

# ─── PDF Generator ──────────────────────────────────────────────────────────

def generate_pdf(data: dict, analysis: str, out_path: Path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

    PRIMARY  = colors.HexColor("#1A3A5C")
    ACCENT   = colors.HexColor("#E63946")
    LIGHT    = colors.HexColor("#EEF2FF")
    MUTED    = colors.HexColor("#64748B")
    BORDER   = colors.HexColor("#E2E8F0")
    BODY_COL = colors.HexColor("#334155")
    WHITE    = colors.white

    company = data.get("company", "Infopercept Consulting Pvt. Ltd.")
    report_type = data.get("report_type", "Business Report")
    period = data.get("period", "FY 2024")

    def header_footer(canvas, doc):
        canvas.saveState()
        w, h = letter
        canvas.setFillColor(PRIMARY)
        canvas.rect(0, h - 0.72*inch, w, 0.72*inch, fill=1, stroke=0)
        canvas.setFillColor(ACCENT)
        canvas.rect(0, h - 0.74*inch, w, 0.025*inch, fill=1, stroke=0)
        canvas.setFillColor(WHITE)
        canvas.setFont("Helvetica-Bold", 14)
        canvas.drawString(0.5*inch, h - 0.38*inch, "INFOPERCEPT")
        canvas.setFont("Helvetica", 10)
        canvas.drawString(1.86*inch, h - 0.38*inch, "Consulting Pvt. Ltd.")
        canvas.setFillColor(colors.HexColor("#B0C8E0"))
        canvas.setFont("Helvetica", 9)
        canvas.drawRightString(w - 0.5*inch, h - 0.35*inch, f"{report_type}  |  {period}")
        canvas.drawRightString(w - 0.5*inch, h - 0.58*inch, company)
        # Footer
        canvas.setFillColor(PRIMARY)
        canvas.rect(0, 0, w, 0.52*inch, fill=1, stroke=0)
        canvas.setFillColor(ACCENT)
        canvas.rect(0, 0.505*inch, w, 0.018*inch, fill=1, stroke=0)
        canvas.setFillColor(colors.HexColor("#B0C8E0"))
        canvas.setFont("Helvetica", 8)
        canvas.drawString(0.5*inch, 0.3*inch,  f"© 2024 {company}  |  Confidential & Proprietary")
        canvas.drawString(0.5*inch, 0.14*inch, f"Generated: {today_str()}")
        canvas.drawRightString(w - 0.5*inch, 0.3*inch,  f"Page {doc.page}")
        canvas.drawRightString(w - 0.5*inch, 0.14*inch, report_type)
        canvas.restoreState()

    doc = SimpleDocTemplate(str(out_path), pagesize=letter,
        rightMargin=0.6*inch, leftMargin=0.6*inch,
        topMargin=1.05*inch, bottomMargin=0.75*inch)

    aw = letter[0] - 1.2*inch   # available width

    S = getSampleStyleSheet()
    title_s = ParagraphStyle("TT", fontSize=26, textColor=PRIMARY,
        fontName="Helvetica-Bold", spaceAfter=6, spaceBefore=10)
    sub_s   = ParagraphStyle("SS", fontSize=13, textColor=ACCENT,
        fontName="Helvetica-Bold", spaceAfter=4)
    meta_s  = ParagraphStyle("MS", fontSize=10, textColor=MUTED, spaceAfter=14)
    h2_s    = ParagraphStyle("H2", fontSize=13, textColor=PRIMARY,
        fontName="Helvetica-Bold", spaceAfter=8, spaceBefore=18)
    body_s  = ParagraphStyle("BS", fontSize=10, textColor=BODY_COL,
        leading=16, spaceAfter=10, alignment=TA_JUSTIFY)
    tbl_h   = ParagraphStyle("TH", fontSize=9, textColor=WHITE,
        fontName="Helvetica-Bold", alignment=TA_CENTER)
    kpi_v   = ParagraphStyle("KV", fontSize=20, textColor=ACCENT,
        fontName="Helvetica-Bold", alignment=TA_CENTER)
    kpi_l   = ParagraphStyle("KL", fontSize=8, textColor=MUTED, alignment=TA_CENTER)
    disc_s  = ParagraphStyle("DS", fontSize=8, textColor=MUTED)

    TABLE_STYLE = TableStyle([
        ("BACKGROUND",   (0,0), (-1,0),  PRIMARY),
        ("TEXTCOLOR",    (0,0), (-1,0),  WHITE),
        ("FONTNAME",     (0,0), (-1,0),  "Helvetica-Bold"),
        ("FONTSIZE",     (0,0), (-1,-1), 9),
        ("ROWBACKGROUNDS",(0,1),(-1,-1), [WHITE, colors.HexColor("#F8FAFC")]),
        ("GRID",         (0,0), (-1,-1), 0.3, BORDER),
        ("TOPPADDING",   (0,0), (-1,-1), 7),
        ("BOTTOMPADDING",(0,0), (-1,-1), 7),
        ("LEFTPADDING",  (0,0), (-1,-1), 8),
        ("RIGHTPADDING", (0,0), (-1,-1), 8),
        ("ALIGN",        (0,0), (-1,0),  "CENTER"),
        ("ALIGN",        (1,1), (-1,-1), "CENTER"),
        ("ALIGN",        (0,1), (0,-1),  "LEFT"),
    ])

    def kpi_grid(items, cols=4):
        cw = aw / cols
        rows = []
        for i in range(0, len(items), cols):
            chunk = items[i:i+cols]
            while len(chunk) < cols: chunk.append(("", ""))
            rows.append([Paragraph(v, kpi_v) for v, _ in chunk])
            rows.append([Paragraph(l, kpi_l) for _, l in chunk])
        t = Table(rows, colWidths=[cw]*cols)
        t.setStyle(TableStyle([
            ("BACKGROUND",  (0,0),(-1,-1), LIGHT),
            ("BOX",         (0,0),(-1,-1), 0.5, BORDER),
            ("INNERGRID",   (0,0),(-1,-1), 0.5, BORDER),
            ("TOPPADDING",  (0,0),(-1,-1), 9),
            ("BOTTOMPADDING",(0,0),(-1,-1),9),
        ]))
        return t

    def make_table(headers, rows, col_ratios=None):
        if col_ratios:
            cws = [aw * r for r in col_ratios]
        else:
            cws = [aw / len(headers)] * len(headers)
        tdata = [[Paragraph(h, tbl_h) for h in headers]] + rows
        t = Table(tdata, colWidths=cws, repeatRows=1)
        t.setStyle(TABLE_STYLE)
        return t

    story = []

    # ── Title block ──
    story.append(Paragraph(report_type, title_s))
    story.append(Paragraph(company, sub_s))
    story.append(Paragraph(
        f"Period: {period}  |  Report Date: {today_str()}  |  Classification: Internal — Confidential",
        meta_s))
    story.append(HRFlowable(width="100%", thickness=2, color=ACCENT, spaceAfter=18))

    # ── KPI grid + tables per report type ──
    if "kpis" in data:                          # ── SALES ──
        k = data["kpis"]
        story.append(kpi_grid([
            (fmt_usd(k["total_revenue"]),        "Total Revenue"),
            (f"{k['attainment_pct']}%",           "Quota Attainment"),
            (f"{k['yoy_growth_pct']}%",           "YoY Growth"),
            (f"{k['win_rate_pct']}%",             "Win Rate"),
            (str(k["total_deals"]),               "Deals Closed"),
            (f"{k['avg_sales_cycle_days']} days", "Avg Sales Cycle"),
            (f"${k['customer_lifetime_value']:,}","Customer LTV"),
            (f"{k['net_revenue_retention_pct']}%","Net Rev. Retention"),
        ]))
        story.append(Spacer(1, 0.2*inch))

        _write_analysis_sections(story, analysis, h2_s, body_s, BORDER, HRFlowable, Spacer, inch)

        story.append(Paragraph("Quarterly Sales Performance", h2_s))
        story.append(HRFlowable(width="100%", thickness=1, color=BORDER, spaceAfter=8))
        story.append(make_table(
            ["Quarter", "Revenue", "Target", "Deals", "Avg Deal Size"],
            [[q["quarter"], fmt_usd(q["revenue"]), fmt_usd(q["target"]),
              str(q["deals_closed"]), fmt_usd(q["avg_deal_size"])]
             for q in data["quarterly_sales"]],
            [0.22, 0.19, 0.19, 0.17, 0.23]
        ))
        story.append(Spacer(1, 0.18*inch))

        story.append(Paragraph("Regional Breakdown", h2_s))
        story.append(HRFlowable(width="100%", thickness=1, color=BORDER, spaceAfter=8))
        story.append(make_table(
            ["Region", "Revenue", "YoY Growth", "Market Share"],
            [[r["region"], fmt_usd(r["revenue"]), f"{r['growth_pct']}%", f"{r['market_share']}%"]
             for r in data["regional_breakdown"]],
            [0.34, 0.22, 0.22, 0.22]
        ))
        story.append(Spacer(1, 0.18*inch))

        story.append(Paragraph("Product Line Performance", h2_s))
        story.append(HRFlowable(width="100%", thickness=1, color=BORDER, spaceAfter=8))
        story.append(make_table(
            ["Product", "Revenue", "Units Sold", "YoY Growth"],
            [[p["product"], fmt_usd(p["revenue"]), str(p["units_sold"]), f"{p['growth_pct']}%"]
             for p in data["top_products"]],
            [0.42, 0.20, 0.18, 0.20]
        ))
        story.append(Spacer(1, 0.18*inch))

        story.append(Paragraph("Sales Team Performance", h2_s))
        story.append(HRFlowable(width="100%", thickness=1, color=BORDER, spaceAfter=8))
        story.append(make_table(
            ["Sales Executive", "Region", "Revenue", "Quota Attainment"],
            [[s["name"], s["region"], fmt_usd(s["revenue"]), f"{s['quota_pct']}%"]
             for s in data["sales_team"]],
            [0.30, 0.28, 0.20, 0.22]
        ))

    elif "headcount" in data:                   # ── HR ──
        h = data["headcount"]
        e = data["engagement"]
        story.append(kpi_grid([
            (str(h["total_employees"]),             "Total Employees"),
            (str(h["new_hires_ytd"]),               "New Hires FY24"),
            (f"{e['voluntary_attrition_pct']}%",    "Attrition Rate"),
            (str(e["eNPS"]),                        "eNPS Score"),
            (f"{e['engagement_score']}/100",        "Engagement Score"),
            (f"{e['training_hours_per_employee']}h","Training / Employee"),
            (f"{data['diversity']['gender_female_pct']}%","Female Workforce"),
            (f"{data['diversity']['leadership_female_pct']}%","Female Leadership"),
        ]))
        story.append(Spacer(1, 0.2*inch))

        _write_analysis_sections(story, analysis, h2_s, body_s, BORDER, HRFlowable, Spacer, inch)

        story.append(Paragraph("Headcount by Department", h2_s))
        story.append(HRFlowable(width="100%", thickness=1, color=BORDER, spaceAfter=8))
        story.append(make_table(
            ["Department", "Headcount", "% of Total", "Avg Tenure (yrs)"],
            [[d["name"], str(d["headcount"]), f"{d['pct']}%", str(d["avg_tenure_years"])]
             for d in data["departments"]],
            [0.40, 0.20, 0.20, 0.20]
        ))
        story.append(Spacer(1, 0.18*inch))

        story.append(Paragraph("Hiring vs. Attrition by Quarter", h2_s))
        story.append(HRFlowable(width="100%", thickness=1, color=BORDER, spaceAfter=8))
        story.append(make_table(
            ["Quarter", "New Hires", "Attrition", "Net Change"],
            [[q["quarter"], str(q["hires"]), str(q["attrition"]),
              f"+{q['hires']-q['attrition']}"]
             for q in data["hiring_by_quarter"]],
            [0.25, 0.25, 0.25, 0.25]
        ))

    elif "income_statement" in data:            # ── FINANCE ──
        IS = data["income_statement"]
        cf = data["cash_flow"]
        story.append(kpi_grid([
            (fmt_usd(IS["total_revenue"]),   "Total Revenue"),
            (f"{IS['gross_margin_pct']}%",   "Gross Margin"),
            (fmt_usd(IS["ebitda"]),          "EBITDA"),
            (f"{IS['ebitda_margin_pct']}%",  "EBITDA Margin"),
            (fmt_usd(IS["net_income"]),      "Net Income"),
            (f"{IS['net_margin_pct']}%",     "Net Margin"),
            (fmt_usd(cf["free_cash_flow"]),  "Free Cash Flow"),
            (fmt_usd(data["balance_sheet"]["cash_equivalents"]), "Cash & Equivalents"),
        ]))
        story.append(Spacer(1, 0.2*inch))

        _write_analysis_sections(story, analysis, h2_s, body_s, BORDER, HRFlowable, Spacer, inch)

        story.append(Paragraph("Income Statement Summary", h2_s))
        story.append(HRFlowable(width="100%", thickness=1, color=BORDER, spaceAfter=8))
        story.append(make_table(
            ["Line Item", "Amount (USD)", "% of Revenue"],
            [
                ["Total Revenue",          fmt_usd(IS["total_revenue"]),  "100.0%"],
                ["Cost of Revenue",        fmt_usd(IS["cost_of_revenue"]),f"{IS['cost_of_revenue']/IS['total_revenue']*100:.1f}%"],
                ["Gross Profit",           fmt_usd(IS["gross_profit"]),   f"{IS['gross_margin_pct']}%"],
                ["Total OpEx",             fmt_usd(IS["operating_expenses"]["total"]),
                 f"{IS['operating_expenses']['total']/IS['total_revenue']*100:.1f}%"],
                ["EBITDA",                 fmt_usd(IS["ebitda"]),         f"{IS['ebitda_margin_pct']}%"],
                ["Net Income",             fmt_usd(IS["net_income"]),     f"{IS['net_margin_pct']}%"],
            ],
            [0.50, 0.28, 0.22]
        ))
        story.append(Spacer(1, 0.18*inch))

        story.append(Paragraph("Revenue by Type", h2_s))
        story.append(HRFlowable(width="100%", thickness=1, color=BORDER, spaceAfter=8))
        story.append(make_table(
            ["Revenue Type", "Amount", "% of Total", "YoY Growth"],
            [[r["type"], fmt_usd(r["amount"]), f"{r['pct']}%", f"{r['yoy_growth']}%"]
             for r in data["revenue_by_type"]],
            [0.34, 0.22, 0.22, 0.22]
        ))
        story.append(Spacer(1, 0.18*inch))

        story.append(Paragraph("Quarterly Financial Performance", h2_s))
        story.append(HRFlowable(width="100%", thickness=1, color=BORDER, spaceAfter=8))
        story.append(make_table(
            ["Quarter", "Revenue", "Gross Profit", "Net Income"],
            [[q["quarter"], fmt_usd(q["revenue"]),
              fmt_usd(q["gross_profit"]), fmt_usd(q["net_income"])]
             for q in data["quarterly_financials"]],
            [0.25, 0.25, 0.25, 0.25]
        ))

    # ── Disclaimer ──
    story.append(Spacer(1, 0.35*inch))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER, spaceAfter=8))
    story.append(Paragraph(
        "<b>Confidentiality Notice:</b> This report is prepared by Infopercept Consulting Pvt. Ltd. "
        "for internal use only. All information is confidential and proprietary. Unauthorized "
        "distribution or reproduction is strictly prohibited. All financial figures in USD unless stated.",
        disc_s))

    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)


def _write_analysis_sections(story, analysis, h2_s, body_s, BORDER, HRFlowable, Spacer, inch):
    """Parse the LLM markdown analysis and add styled paragraphs to the story."""
    if not analysis:
        return
    current_heading = None
    for line in analysis.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            story.append(Spacer(1, 0.05*inch))
            story.append(HRFlowable(width="100%", thickness=1, color=BORDER, spaceAfter=6))
            from reportlab.platypus import Paragraph as P
            story.append(P(heading, h2_s))
            current_heading = heading
        else:
            from reportlab.platypus import Paragraph as P
            story.append(P(line, body_s))
    story.append(Spacer(1, 0.1*inch))


# ─── DOCX Generator ─────────────────────────────────────────────────────────

def generate_docx(data: dict, analysis: str, out_path: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    PRIMARY_RGB  = RGBColor(0x1A, 0x3A, 0x5C)
    ACCENT_RGB   = RGBColor(0xE6, 0x39, 0x46)
    MUTED_RGB    = RGBColor(0x64, 0x74, 0x8B)
    BODY_RGB     = RGBColor(0x33, 0x41, 0x55)
    HEADER_FILL  = "1A3A5C"
    ALT_ROW      = "F8FAFC"
    WHITE_FILL   = "FFFFFF"
    BORDER_CLR   = "CCCCCC"

    company     = data.get("company", "Infopercept Consulting Pvt. Ltd.")
    report_type = data.get("report_type", "Business Report")
    period      = data.get("period", "FY 2024")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    def set_cell_shading(cell, fill_hex):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex)
        tcPr.append(shd)

    def set_cell_border(cell):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ["top", "left", "bottom", "right"]:
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"),   "single")
            b.set(qn("w:sz"),    "4")
            b.set(qn("w:color"), BORDER_CLR)
            tcBorders.append(b)
        tcPr.append(tcBorders)

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(8)
        run = p.add_run(text)
        run.bold      = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.color.rgb = PRIMARY_RGB
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "8")
        bottom.set(qn("w:color"), "E63946")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_body(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_RGB
        return p

    def add_kpi_table(items, cols=4):
        while len(items) % cols != 0:
            items.append(("", ""))
        rows = len(items) // cols
        t = doc.add_table(rows=rows*2, cols=cols)
        t.style = "Table Grid"
        idx = 0
        for r in range(rows):
            for c_i in range(cols):
                val_cell = t.cell(r*2,   c_i)
                lbl_cell = t.cell(r*2+1, c_i)
                v, l = items[idx]; idx += 1
                set_cell_shading(val_cell, "EEF2FF")
                set_cell_shading(lbl_cell, "EEF2FF")
                set_cell_border(val_cell)
                set_cell_border(lbl_cell)
                vp = val_cell.paragraphs[0]
                vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                vr = vp.add_run(v)
                vr.bold = True; vr.font.size = Pt(18); vr.font.color.rgb = ACCENT_RGB
                lp = lbl_cell.paragraphs[0]
                lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lr = lp.add_run(l)
                lr.font.size = Pt(8); lr.font.color.rgb = MUTED_RGB
        doc.add_paragraph()

    def add_data_table(headers, rows):
        t = doc.add_table(rows=1+len(rows), cols=len(headers))
        t.style = "Table Grid"
        # Header row
        hrow = t.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            set_cell_shading(cell, HEADER_FILL)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True; run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Data rows
        for ri, row in enumerate(rows):
            trow = t.rows[ri+1]
            fill = WHITE_FILL if ri % 2 == 0 else ALT_ROW
            for ci, cell_val in enumerate(row):
                cell = trow.cells[ci]
                set_cell_shading(cell, fill)
                set_cell_border(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(str(cell_val))
                run.font.size = Pt(9)
                run.font.color.rgb = BODY_RGB
        doc.add_paragraph()

    # ── Header (simple paragraph, since docx header styling is complex) ──
    hdr_sec = doc.sections[0]
    hdr = hdr_sec.header
    hp = hdr.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = hp.add_run("INFOPERCEPT")
    r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = PRIMARY_RGB
    r2 = hp.add_run(f"  |  Consulting Pvt. Ltd.      {report_type}  |  {period}")
    r2.font.size = Pt(9); r2.font.color.rgb = MUTED_RGB
    # Header border
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "12"); bot.set(qn("w:color"), HEADER_FILL)
    pBdr.append(bot); pPr.append(pBdr)

    # Footer
    ftr = hdr_sec.footer
    fp = ftr.paragraphs[0]
    fp.clear()
    fr1 = fp.add_run(f"© 2024 {company}  |  Confidential & Proprietary     Generated: {today_str()}")
    fr1.font.size = Pt(8); fr1.font.color.rgb = MUTED_RGB
    pPr2 = fp._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top2  = OxmlElement("w:top")
    top2.set(qn("w:val"), "single"); top2.set(qn("w:sz"), "4"); top2.set(qn("w:color"), BORDER_CLR)
    pBdr2.append(top2); pPr2.append(pBdr2)

    # ── Title block ──
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after  = Pt(4)
    tr = tp.add_run(report_type)
    tr.bold = True; tr.font.size = Pt(24); tr.font.color.rgb = PRIMARY_RGB

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(3)
    sr = sp.add_run(company)
    sr.bold = True; sr.font.size = Pt(13); sr.font.color.rgb = ACCENT_RGB

    mp = doc.add_paragraph()
    mp.paragraph_format.space_after = Pt(12)
    mr = mp.add_run(f"Period: {period}  |  Report Date: {today_str()}  |  Classification: Internal — Confidential")
    mr.font.size = Pt(9); mr.font.color.rgb = MUTED_RGB

    # ── KPI + data tables ──
    if "kpis" in data:
        k = data["kpis"]
        add_kpi_table([
            (fmt_usd(k["total_revenue"]),        "Total Revenue"),
            (f"{k['attainment_pct']}%",          "Quota Attainment"),
            (f"{k['yoy_growth_pct']}%",          "YoY Growth"),
            (f"{k['win_rate_pct']}%",            "Win Rate"),
            (str(k["total_deals"]),              "Deals Closed"),
            (f"{k['avg_sales_cycle_days']} days","Avg Sales Cycle"),
            (f"${k['customer_lifetime_value']:,}","Customer LTV"),
            (f"{k['net_revenue_retention_pct']}%","Net Rev. Retention"),
        ])
        _write_docx_analysis(doc, analysis, add_heading, add_body)

        add_heading("Quarterly Sales Performance")
        add_data_table(
            ["Quarter", "Revenue", "Target", "Deals", "Avg Deal Size"],
            [[q["quarter"], fmt_usd(q["revenue"]), fmt_usd(q["target"]),
              str(q["deals_closed"]), fmt_usd(q["avg_deal_size"])]
             for q in data["quarterly_sales"]]
        )
        add_heading("Regional Breakdown")
        add_data_table(
            ["Region", "Revenue", "YoY Growth", "Market Share"],
            [[r["region"], fmt_usd(r["revenue"]),
              f"{r['growth_pct']}%", f"{r['market_share']}%"]
             for r in data["regional_breakdown"]]
        )
        add_heading("Product Line Performance")
        add_data_table(
            ["Product", "Revenue", "Units Sold", "YoY Growth"],
            [[p["product"], fmt_usd(p["revenue"]),
              str(p["units_sold"]), f"{p['growth_pct']}%"]
             for p in data["top_products"]]
        )
        add_heading("Sales Team Performance")
        add_data_table(
            ["Sales Executive", "Region", "Revenue", "Quota Attainment"],
            [[s["name"], s["region"], fmt_usd(s["revenue"]),
              f"{s['quota_pct']}%"] for s in data["sales_team"]]
        )

    elif "headcount" in data:
        h = data["headcount"]; e = data["engagement"]
        add_kpi_table([
            (str(h["total_employees"]),             "Total Employees"),
            (str(h["new_hires_ytd"]),               "New Hires FY24"),
            (f"{e['voluntary_attrition_pct']}%",   "Attrition Rate"),
            (str(e["eNPS"]),                        "eNPS Score"),
            (f"{e['engagement_score']}/100",        "Engagement Score"),
            (f"{e['training_hours_per_employee']}h","Training / Employee"),
            (f"{data['diversity']['gender_female_pct']}%","Female Workforce"),
            (f"{data['diversity']['leadership_female_pct']}%","Female Leadership"),
        ])
        _write_docx_analysis(doc, analysis, add_heading, add_body)
        add_heading("Headcount by Department")
        add_data_table(
            ["Department", "Headcount", "% of Total", "Avg Tenure (yrs)"],
            [[d["name"], str(d["headcount"]), f"{d['pct']}%", str(d["avg_tenure_years"])]
             for d in data["departments"]]
        )
        add_heading("Hiring vs. Attrition by Quarter")
        add_data_table(
            ["Quarter", "New Hires", "Attrition", "Net Change"],
            [[q["quarter"], str(q["hires"]), str(q["attrition"]),
              f"+{q['hires']-q['attrition']}"] for q in data["hiring_by_quarter"]]
        )

    elif "income_statement" in data:
        IS = data["income_statement"]; cf = data["cash_flow"]
        add_kpi_table([
            (fmt_usd(IS["total_revenue"]),  "Total Revenue"),
            (f"{IS['gross_margin_pct']}%",  "Gross Margin"),
            (fmt_usd(IS["ebitda"]),         "EBITDA"),
            (f"{IS['ebitda_margin_pct']}%", "EBITDA Margin"),
            (fmt_usd(IS["net_income"]),     "Net Income"),
            (f"{IS['net_margin_pct']}%",    "Net Margin"),
            (fmt_usd(cf["free_cash_flow"]), "Free Cash Flow"),
            (fmt_usd(data["balance_sheet"]["cash_equivalents"]),"Cash & Equivalents"),
        ])
        _write_docx_analysis(doc, analysis, add_heading, add_body)
        add_heading("Income Statement Summary")
        add_data_table(
            ["Line Item", "Amount (USD)", "% of Revenue"],
            [
                ["Total Revenue",   fmt_usd(IS["total_revenue"]),  "100.0%"],
                ["Cost of Revenue", fmt_usd(IS["cost_of_revenue"]),
                 f"{IS['cost_of_revenue']/IS['total_revenue']*100:.1f}%"],
                ["Gross Profit",    fmt_usd(IS["gross_profit"]),   f"{IS['gross_margin_pct']}%"],
                ["Total OpEx",      fmt_usd(IS["operating_expenses"]["total"]),
                 f"{IS['operating_expenses']['total']/IS['total_revenue']*100:.1f}%"],
                ["EBITDA",          fmt_usd(IS["ebitda"]),         f"{IS['ebitda_margin_pct']}%"],
                ["Net Income",      fmt_usd(IS["net_income"]),     f"{IS['net_margin_pct']}%"],
            ]
        )
        add_heading("Revenue by Type")
        add_data_table(
            ["Revenue Type", "Amount", "% of Total", "YoY Growth"],
            [[r["type"], fmt_usd(r["amount"]),
              f"{r['pct']}%", f"{r['yoy_growth']}%"]
             for r in data["revenue_by_type"]]
        )

    # Disclaimer
    dp = doc.add_paragraph()
    dp.paragraph_format.space_before = Pt(24)
    dr = dp.add_run(
        "Confidentiality Notice: This report is prepared by Infopercept Consulting Pvt. Ltd. "
        "for internal use only. Unauthorized distribution is prohibited.")
    dr.font.size = Pt(8); dr.font.color.rgb = MUTED_RGB

    doc.save(str(out_path))


def _write_docx_analysis(doc, analysis, add_heading, add_body):
    if not analysis:
        return
    for line in analysis.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            add_heading(line[3:].strip())
        else:
            add_body(line)


# ─── PPTX Generator ─────────────────────────────────────────────────────────

def generate_pptx(data: dict, analysis: str, out_path: Path):
    """Generate PPTX by calling the Node.js pptxgenjs helper."""
    import subprocess, tempfile

    helper = BASE_DIR / "_pptx_helper.js"
    payload = {
        "data": data,
        "analysis": analysis,
        "outPath": str(out_path),
        "today": today_str(),
    }
    payload_file = Path(tempfile.mktemp(suffix=".json"))
    payload_file.write_text(json.dumps(payload))

    if not helper.exists():
        _write_pptx_helper(helper)

    result = subprocess.run(
        ["node", str(helper), str(payload_file)],
        capture_output=True, text=True, timeout=60
    )
    payload_file.unlink(missing_ok=True)
    if result.returncode != 0:
        raise RuntimeError(f"PPTX helper failed:\n{result.stderr}")


def _write_pptx_helper(path: Path):
    """Write the Node.js PPTX helper script."""
    path.write_text(r"""
const fs = require('fs');
const pptxgen = require('pptxgenjs');

const payload = JSON.parse(fs.readFileSync(process.argv[2], 'utf8'));
const { data, analysis, outPath, today } = payload;

async function run() {
const pres = new pptxgen();
pres.layout = 'LAYOUT_16x9';
pres.author = 'Infopercept Report System';
pres.title  = data.report_type || 'Report';

const P   = '1A3A5C', A = 'E63946', W = 'FFFFFF', BODY = '334155';
const MUTED = '94A3B8', GOLD = 'F4A261', DARK = '0F2240', LT = 'F0F4F8';
const company     = data.company     || 'Infopercept Consulting Pvt. Ltd.';
const report_type = data.report_type || 'Business Report';
const period      = data.period      || 'FY 2024';

function fmt(n){ return '$'+(n/1e6).toFixed(2)+'M'; }

function chrome(slide, title) {
  slide.addShape(pres.shapes.RECTANGLE,{x:0,y:0,w:10,h:0.62,fill:{color:P},line:{color:P}});
  slide.addShape(pres.shapes.RECTANGLE,{x:0,y:0.60,w:10,h:0.04,fill:{color:A},line:{color:A}});
  slide.addText('INFOPERCEPT',{x:0.28,y:0.06,w:2.4,h:0.28,fontSize:13,bold:true,color:W,fontFace:'Arial',margin:0});
  slide.addText('Consulting Pvt. Ltd.',{x:0.28,y:0.33,w:3,h:0.2,fontSize:8.5,color:'B0C8E0',fontFace:'Arial',margin:0});
  slide.addText(title,{x:3.5,y:0.15,w:6.2,h:0.3,fontSize:9,color:'B0C8E0',fontFace:'Arial',align:'right',margin:0});
  slide.addShape(pres.shapes.RECTANGLE,{x:0,y:5.325,w:10,h:0.3,fill:{color:P},line:{color:P}});
  slide.addShape(pres.shapes.RECTANGLE,{x:0,y:5.322,w:10,h:0.014,fill:{color:A},line:{color:A}});
  slide.addText('© 2024 '+company+'  |  Confidential & Proprietary',{x:0.28,y:5.345,w:6,h:0.24,fontSize:7.5,color:'B0C8E0',fontFace:'Arial',margin:0});
  slide.addText(period+'  |  '+today,{x:6.5,y:5.345,w:3.2,h:0.24,fontSize:7.5,color:'B0C8E0',fontFace:'Arial',align:'right',margin:0});
}

// ── Slide 1: Title ──
const s1 = pres.addSlide();
s1.addShape(pres.shapes.RECTANGLE,{x:0,y:0,w:10,h:5.625,fill:{color:P},line:{color:P}});
s1.addShape(pres.shapes.RECTANGLE,{x:0,y:3.65,w:10,h:1.975,fill:{color:DARK},line:{color:DARK}});
s1.addShape(pres.shapes.RECTANGLE,{x:0,y:0,w:10,h:0.07,fill:{color:A},line:{color:A}});
s1.addShape(pres.shapes.RECTANGLE,{x:0,y:5.555,w:10,h:0.07,fill:{color:A},line:{color:A}});
s1.addShape(pres.shapes.RECTANGLE,{x:0.45,y:1.25,w:0.1,h:2.1,fill:{color:A},line:{color:A}});
s1.addText('INFOPERCEPT',{x:0.7,y:1.25,w:9,h:0.4,fontSize:11,bold:true,color:MUTED,fontFace:'Arial',charSpacing:6,margin:0});
s1.addText(report_type,{x:0.7,y:1.72,w:8.8,h:0.95,fontSize:36,bold:true,color:W,fontFace:'Arial',margin:0});
s1.addText(period,{x:0.7,y:2.74,w:8,h:0.38,fontSize:14,color:'B0C8E0',fontFace:'Arial',margin:0});
s1.addText(company,{x:0.55,y:3.76,w:9,h:0.36,fontSize:12,bold:true,color:GOLD,fontFace:'Arial',margin:0});
s1.addText('Report Date: '+today+'  |  Internal — Confidential',{x:0.55,y:4.22,w:9,h:0.27,fontSize:9,color:MUTED,fontFace:'Arial',margin:0});

// ── Slide 2: KPI Cards ──
let kpis = [];
if (data.kpis) {
  const k = data.kpis;
  kpis = [
    {v:fmt(k.total_revenue),l:'Total Revenue'},
    {v:k.attainment_pct+'%',l:'Quota Attainment'},
    {v:k.yoy_growth_pct+'%',l:'YoY Growth'},
    {v:k.win_rate_pct+'%',l:'Win Rate'},
    {v:String(k.total_deals),l:'Deals Closed'},
    {v:k.net_revenue_retention_pct+'%',l:'Net Rev. Retention'},
  ];
} else if (data.headcount) {
  const h=data.headcount,e=data.engagement;
  kpis=[
    {v:String(h.total_employees),l:'Total Employees'},
    {v:String(h.new_hires_ytd),l:'New Hires FY24'},
    {v:e.voluntary_attrition_pct+'%',l:'Attrition Rate'},
    {v:String(e.eNPS),l:'eNPS Score'},
    {v:e.engagement_score+'/100',l:'Engagement Score'},
    {v:data.diversity.gender_female_pct+'%',l:'Female Workforce'},
  ];
} else if (data.income_statement) {
  const IS=data.income_statement,cf=data.cash_flow;
  kpis=[
    {v:fmt(IS.total_revenue),l:'Total Revenue'},
    {v:IS.gross_margin_pct+'%',l:'Gross Margin'},
    {v:fmt(IS.ebitda),l:'EBITDA'},
    {v:IS.ebitda_margin_pct+'%',l:'EBITDA Margin'},
    {v:fmt(IS.net_income),l:'Net Income'},
    {v:IS.net_margin_pct+'%',l:'Net Margin'},
  ];
}

const s2 = pres.addSlide();
chrome(s2,'Key Performance Indicators');
s2.addShape(pres.shapes.RECTANGLE,{x:0,y:0.64,w:10,h:4.685,fill:{color:LT},line:{color:LT}});
s2.addText('Key Performance Indicators — '+period,{x:0.4,y:0.76,w:9,h:0.42,fontSize:19,bold:true,color:P,fontFace:'Arial',margin:0});
s2.addShape(pres.shapes.LINE,{x:0.4,y:1.18,w:9.2,h:0,line:{color:'DDDDDD',width:1}});
kpis.forEach((kpi,i)=>{
  const col=i%3, row=Math.floor(i/3);
  const x=0.38+col*3.1, y=1.28+row*1.72;
  s2.addShape(pres.shapes.RECTANGLE,{x,y,w:2.92,h:1.56,fill:{color:W},line:{color:'E2E8F0',width:1},
    shadow:{type:'outer',blur:5,offset:2,angle:135,color:'000000',opacity:0.06}});
  s2.addShape(pres.shapes.RECTANGLE,{x,y,w:0.07,h:1.56,fill:{color:i===0?A:P},line:{color:i===0?A:P}});
  s2.addText(kpi.v,{x:x+0.12,y:y+0.28,w:2.72,h:0.65,fontSize:28,bold:true,color:A,fontFace:'Arial',align:'center',margin:0});
  s2.addText(kpi.l,{x:x+0.12,y:y+1.02,w:2.72,h:0.38,fontSize:10,color:BODY,fontFace:'Arial',align:'center',margin:0});
});

// ── Slide 3: Chart ──
const s3 = pres.addSlide();
chrome(s3,'Performance Trend');
s3.addShape(pres.shapes.RECTANGLE,{x:0,y:0.64,w:10,h:4.685,fill:{color:LT},line:{color:LT}});
s3.addText('Quarterly Performance Trend',{x:0.4,y:0.76,w:9,h:0.42,fontSize:19,bold:true,color:P,fontFace:'Arial',margin:0});

let chartLabels=[], chartVals=[], chartVals2=[];
if (data.quarterly_sales) {
  chartLabels = data.quarterly_sales.map(q=>q.quarter);
  chartVals   = data.quarterly_sales.map(q=>parseFloat((q.revenue/1e6).toFixed(2)));
  chartVals2  = data.quarterly_sales.map(q=>parseFloat((q.target/1e6).toFixed(2)));
  s3.addChart(pres.charts.BAR,[
    {name:'Revenue ($M)',labels:chartLabels,values:chartVals},
    {name:'Target ($M)', labels:chartLabels,values:chartVals2}
  ],{x:0.4,y:1.25,w:9.2,h:3.7,barDir:'col',barGrouping:'clustered',
    chartColors:[P,A],chartArea:{fill:{color:W},roundedCorners:true},
    catAxisLabelColor:BODY,valAxisLabelColor:BODY,
    valGridLine:{color:'E2E8F0',size:0.5},catGridLine:{style:'none'},
    showValue:true,dataLabelFontSize:9,showLegend:true,legendPos:'b',legendFontSize:10,showTitle:false});
} else if (data.hiring_by_quarter) {
  chartLabels = data.hiring_by_quarter.map(q=>q.quarter);
  chartVals   = data.hiring_by_quarter.map(q=>q.hires);
  chartVals2  = data.hiring_by_quarter.map(q=>q.attrition);
  s3.addChart(pres.charts.BAR,[
    {name:'New Hires',labels:chartLabels,values:chartVals},
    {name:'Attrition',labels:chartLabels,values:chartVals2}
  ],{x:0.4,y:1.25,w:9.2,h:3.7,barDir:'col',barGrouping:'clustered',
    chartColors:[P,A],chartArea:{fill:{color:W},roundedCorners:true},
    catAxisLabelColor:BODY,valAxisLabelColor:BODY,
    valGridLine:{color:'E2E8F0',size:0.5},catGridLine:{style:'none'},
    showValue:true,dataLabelFontSize:9,showLegend:true,legendPos:'b',legendFontSize:10,showTitle:false});
} else if (data.quarterly_financials) {
  chartLabels = data.quarterly_financials.map(q=>q.quarter);
  chartVals   = data.quarterly_financials.map(q=>parseFloat((q.revenue/1e6).toFixed(2)));
  chartVals2  = data.quarterly_financials.map(q=>parseFloat((q.net_income/1e6).toFixed(2)));
  s3.addChart(pres.charts.BAR,[
    {name:'Revenue ($M)',  labels:chartLabels,values:chartVals},
    {name:'Net Income ($M)',labels:chartLabels,values:chartVals2}
  ],{x:0.4,y:1.25,w:9.2,h:3.7,barDir:'col',barGrouping:'clustered',
    chartColors:[P,A],chartArea:{fill:{color:W},roundedCorners:true},
    catAxisLabelColor:BODY,valAxisLabelColor:BODY,
    valGridLine:{color:'E2E8F0',size:0.5},catGridLine:{style:'none'},
    showValue:true,dataLabelFontSize:9,showLegend:true,legendPos:'b',legendFontSize:10,showTitle:false});
}

// ── Slide 4: Analysis ──
const s4 = pres.addSlide();
chrome(s4,'Analysis & Insights');
s4.addShape(pres.shapes.RECTANGLE,{x:0,y:0.64,w:10,h:4.685,fill:{color:LT},line:{color:LT}});
s4.addText('Analysis & Key Insights',{x:0.4,y:0.76,w:9,h:0.42,fontSize:19,bold:true,color:P,fontFace:'Arial',margin:0});
// Parse analysis sections
const lines = (analysis||'').split('\n').filter(l=>l.trim());
let yPos = 1.28;
for (const line of lines.slice(0,12)) {
  if (line.startsWith('## ')) {
    s4.addText(line.slice(3).trim(),{x:0.4,y:yPos,w:9.2,h:0.32,fontSize:12,bold:true,color:P,fontFace:'Arial',margin:0});
    yPos += 0.34;
  } else if (yPos < 4.8) {
    s4.addText(line.trim(),{x:0.5,y:yPos,w:9.0,h:0.32,fontSize:10,color:BODY,fontFace:'Arial',wrap:true,margin:0});
    yPos += 0.36;
  }
}

// ── Slide 5: Data Table ──
const s5 = pres.addSlide();
let tableTitle='', tableRows=[];
if (data.quarterly_sales) {
  tableTitle='Quarterly Sales Performance';
  tableRows=[
    [{text:'Quarter',options:{fill:{color:P},color:W,bold:true,align:'center'}},
     {text:'Revenue',options:{fill:{color:P},color:W,bold:true,align:'center'}},
     {text:'Target',options:{fill:{color:P},color:W,bold:true,align:'center'}},
     {text:'Deals Closed',options:{fill:{color:P},color:W,bold:true,align:'center'}},
     {text:'Avg Deal Size',options:{fill:{color:P},color:W,bold:true,align:'center'}}],
    ...data.quarterly_sales.map((q,i)=>[
      {text:q.quarter,options:{fill:{color:i%2===0?W:'F8FAFC'},align:'left'}},
      {text:fmt(q.revenue),options:{fill:{color:i%2===0?W:'F8FAFC'},align:'center'}},
      {text:fmt(q.target),options:{fill:{color:i%2===0?W:'F8FAFC'},align:'center'}},
      {text:String(q.deals_closed),options:{fill:{color:i%2===0?W:'F8FAFC'},align:'center'}},
      {text:'$'+q.avg_deal_size.toLocaleString(),options:{fill:{color:i%2===0?W:'F8FAFC'},align:'center'}},
    ])
  ];
} else if (data.departments) {
  tableTitle='Headcount by Department';
  tableRows=[
    [{text:'Department',options:{fill:{color:P},color:W,bold:true}},
     {text:'Headcount',options:{fill:{color:P},color:W,bold:true,align:'center'}},
     {text:'% of Total',options:{fill:{color:P},color:W,bold:true,align:'center'}},
     {text:'Avg Tenure (yrs)',options:{fill:{color:P},color:W,bold:true,align:'center'}}],
    ...data.departments.map((d,i)=>[
      {text:d.name,options:{fill:{color:i%2===0?W:'F8FAFC'}}},
      {text:String(d.headcount),options:{fill:{color:i%2===0?W:'F8FAFC'},align:'center'}},
      {text:d.pct+'%',options:{fill:{color:i%2===0?W:'F8FAFC'},align:'center'}},
      {text:String(d.avg_tenure_years),options:{fill:{color:i%2===0?W:'F8FAFC'},align:'center'}},
    ])
  ];
} else if (data.income_statement) {
  tableTitle='Income Statement Summary';
  const IS=data.income_statement;
  tableRows=[
    [{text:'Line Item',options:{fill:{color:P},color:W,bold:true}},
     {text:'Amount',options:{fill:{color:P},color:W,bold:true,align:'center'}},
     {text:'% of Revenue',options:{fill:{color:P},color:W,bold:true,align:'center'}}],
    ...([
      ['Total Revenue',fmt(IS.total_revenue),'100%'],
      ['Cost of Revenue',fmt(IS.cost_of_revenue),(IS.cost_of_revenue/IS.total_revenue*100).toFixed(1)+'%'],
      ['Gross Profit',fmt(IS.gross_profit),IS.gross_margin_pct+'%'],
      ['Total OpEx',fmt(IS.operating_expenses.total),(IS.operating_expenses.total/IS.total_revenue*100).toFixed(1)+'%'],
      ['EBITDA',fmt(IS.ebitda),IS.ebitda_margin_pct+'%'],
      ['Net Income',fmt(IS.net_income),IS.net_margin_pct+'%'],
    ].map(([l,v,p],i)=>[
      {text:l,options:{fill:{color:i%2===0?W:'F8FAFC'}}},
      {text:v,options:{fill:{color:i%2===0?W:'F8FAFC'},align:'center'}},
      {text:p,options:{fill:{color:i%2===0?W:'F8FAFC'},align:'center'}},
    ]))
  ];
}
chrome(s5, tableTitle);
s5.addShape(pres.shapes.RECTANGLE,{x:0,y:0.64,w:10,h:4.685,fill:{color:LT},line:{color:LT}});
s5.addText(tableTitle,{x:0.4,y:0.76,w:9,h:0.42,fontSize:19,bold:true,color:P,fontFace:'Arial',margin:0});
if (tableRows.length>0) {
  s5.addTable(tableRows,{x:0.4,y:1.3,w:9.2,border:{pt:0.5,color:'E2E8F0'},fontSize:11,fontFace:'Arial',color:BODY});
}

// ── Slide 6: Recommendations ──
const s6 = pres.addSlide();
chrome(s6,'Strategic Recommendations');
s6.addShape(pres.shapes.RECTANGLE,{x:0,y:0.64,w:10,h:4.685,fill:{color:LT},line:{color:LT}});
s6.addText('Strategic Recommendations',{x:0.4,y:0.76,w:9,h:0.42,fontSize:19,bold:true,color:P,fontFace:'Arial',margin:0});

// Extract recommendations from analysis
const recLines = (analysis||'').split('\n')
  .filter(l=>l.trim() && !l.startsWith('#'));
const recs = recLines.slice(-5).filter(l=>l.length>20);

recs.slice(0,5).forEach((rec,i)=>{
  const y=1.32+i*0.71;
  s6.addShape(pres.shapes.RECTANGLE,{x:0.38,y,w:9.24,h:0.62,fill:{color:W},line:{color:'E2E8F0',width:1}});
  s6.addShape(pres.shapes.RECTANGLE,{x:0.38,y,w:0.58,h:0.62,fill:{color:i===0?A:P},line:{color:i===0?A:P}});
  s6.addText(String(i+1),{x:0.38,y:y+0.08,w:0.58,h:0.46,fontSize:18,bold:true,color:W,fontFace:'Arial',align:'center',margin:0});
  s6.addText(rec.replace(/^\d+\.\s*/,''),{x:1.06,y:y+0.1,w:8.4,h:0.45,fontSize:10.5,color:BODY,fontFace:'Arial',margin:0});
});

// ── Slide 7: Thank You ──
const s7 = pres.addSlide();
s7.addShape(pres.shapes.RECTANGLE,{x:0,y:0,w:10,h:5.625,fill:{color:P},line:{color:P}});
s7.addShape(pres.shapes.RECTANGLE,{x:0,y:0,w:10,h:0.09,fill:{color:A},line:{color:A}});
s7.addShape(pres.shapes.RECTANGLE,{x:0,y:5.535,w:10,h:0.09,fill:{color:A},line:{color:A}});
s7.addShape(pres.shapes.OVAL,{x:3.5,y:0.55,w:3,h:3,fill:{color:'132C47'},line:{color:'1E4A7A',width:1}});
s7.addText('IP',{x:3.5,y:1.5,w:3,h:1.1,fontSize:56,bold:true,color:W,fontFace:'Arial',align:'center',margin:0});
s7.addText('Thank You',{x:0.5,y:3.72,w:9,h:0.72,fontSize:36,bold:true,color:W,fontFace:'Arial',align:'center',margin:0});
s7.addShape(pres.shapes.LINE,{x:2.5,y:4.55,w:5,h:0,line:{color:A,width:2}});
s7.addText(company,{x:0.5,y:4.65,w:9,h:0.32,fontSize:12,color:'B0C8E0',fontFace:'Arial',align:'center',margin:0});
s7.addText('www.infopercept.com  |  contact@infopercept.com  |  +91 79 4896 0000',{x:0.5,y:5.04,w:9,h:0.26,fontSize:9.5,color:MUTED,fontFace:'Arial',align:'center',margin:0});

await pres.writeFile({fileName: outPath});
console.log('OK');
}
run().catch(e=>{ console.error(e.message); process.exit(1); });
""")


# ─── Format selection prompt ─────────────────────────────────────────────────

def ask_output_format() -> str:
    print()
    divider()
    print(c("  What output format would you like?", BOLD, WHITE))
    print()
    print(f"    {c('[1]', CYAN, BOLD)}  📝  Word Document  {c('(.docx)', DIM)}")
    print(f"    {c('[2]', CYAN, BOLD)}  📄  PDF Document   {c('(.pdf)', DIM)}")
    print(f"    {c('[3]', CYAN, BOLD)}  📊  PowerPoint     {c('(.pptx)', DIM)}")
    print(f"    {c('[4]', CYAN, BOLD)}  🗂️   All three formats")
    print()
    divider()

    choices = {"1": "docx", "2": "pdf", "3": "pptx", "4": "all",
               "docx": "docx", "pdf": "pdf", "pptx": "pptx", "word": "docx",
               "powerpoint": "pptx", "all": "all"}

    while True:
        raw = input(c("  ➜  Your choice: ", YELLOW, BOLD)).strip().lower()
        if raw in choices:
            return choices[raw]
        warn(f"Please enter 1, 2, 3, or 4  (got: '{raw}')")


# ─── Main flow ───────────────────────────────────────────────────────────────

def main():
    banner()

    # ── Get the user's prompt ──
    if len(sys.argv) > 1:
        prompt = " ".join(sys.argv[1:])
        info(f"Prompt received: {c(prompt, BOLD)}")
    else:
        print(c("  Type your request below.", DIM))
        print(c("  Examples:  'create a sales report'  |  'generate HR report'  |  'finance summary'\n", DIM))
        prompt = input(c("  ➜  What report do you need?  ", YELLOW, BOLD)).strip()
        if not prompt:
            err("No input provided. Exiting.")
            sys.exit(1)

    print()
    divider()

    # ── Detect which dataset to use ──
    step("🔍", "Analysing your request…")
    report_type = detect_report_type(prompt)

    if not report_type:
        print()
        warn("Could not auto-detect report type from your prompt.")
        print()
        for i, (key, meta) in enumerate(DATA_CATALOGUE.items(), 1):
            print(f"    {c(f'[{i}]', CYAN, BOLD)}  {meta['emoji']}  {meta['label']}")
        print()
        while True:
            choice = input(c("  ➜  Select data source (1/2/3): ", YELLOW, BOLD)).strip()
            keys = list(DATA_CATALOGUE.keys())
            if choice.isdigit() and 1 <= int(choice) <= len(keys):
                report_type = keys[int(choice)-1]
                break
            warn("Please enter 1, 2, or 3.")

    meta = DATA_CATALOGUE[report_type]
    ok(f"Matched data source: {meta['emoji']}  {c(meta['label'], BOLD)}")

    # ── Load data ──
    step("📂", f"Loading {meta['label']} data…")
    try:
        data = load_data(report_type)
    except FileNotFoundError:
        err(f"Data file not found: {meta['file']}")
        sys.exit(1)

    ok(f"Loaded  {c(meta['file'].name, BOLD)}  ({len(json.dumps(data))} bytes)")

    # Show a quick data preview
    print()
    print(c("  ┌─ Data Preview " + "─"*44, DIM))
    if "kpis" in data:
        k = data["kpis"]
        print(c(f"  │  Revenue: {fmt_usd(k['total_revenue'])}  |  Growth: {k['yoy_growth_pct']}%  |  Deals: {k['total_deals']}  |  Win Rate: {k['win_rate_pct']}%", DIM))
    elif "headcount" in data:
        h = data["headcount"]; e = data["engagement"]
        print(c(f"  │  Employees: {h['total_employees']}  |  New Hires: {h['new_hires_ytd']}  |  eNPS: {e['eNPS']}  |  Attrition: {e['voluntary_attrition_pct']}%", DIM))
    elif "income_statement" in data:
        IS = data["income_statement"]
        print(c(f"  │  Revenue: {fmt_usd(IS['total_revenue'])}  |  Gross Margin: {IS['gross_margin_pct']}%  |  EBITDA: {fmt_usd(IS['ebitda'])}  |  Net Income: {fmt_usd(IS['net_income'])}", DIM))
    print(c("  └" + "─"*59, DIM))

    # ── Ask output format ──
    fmt_choice = ask_output_format()
    formats = ["docx", "pdf", "pptx"] if fmt_choice == "all" else [fmt_choice]

    # ── Generate AI analysis ──
    print()
    divider()
    api_key = get_api_key()
    analysis = ""

    if api_key:
        step("🤖", "Generating AI analysis via Claude…")
        try:
            analysis = generate_analysis(data, prompt, api_key)
            ok("AI analysis complete.")
        except Exception as ex:
            warn(f"AI analysis skipped ({ex}). Using structured data only.")
    else:
        warn("ANTHROPIC_API_KEY not set — skipping AI narrative. Set it for richer reports.")
        info("Export:  export ANTHROPIC_API_KEY=sk-ant-...")

    # ── Build report(s) ──
    print()
    divider()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"Infopercept-{meta['label'].replace(' ','-')}-Report_{timestamp}"
    generated = []

    for fmt in formats:
        out_path = OUT_DIR / f"{base_name}.{fmt}"
        step({"docx":"📝","pdf":"📄","pptx":"📊"}[fmt],
             f"Building {c(fmt.upper(), BOLD, WHITE)} report…")
        try:
            if fmt == "pdf":
                generate_pdf(data, analysis, out_path)
            elif fmt == "docx":
                generate_docx(data, analysis, out_path)
            elif fmt == "pptx":
                generate_pptx(data, analysis, out_path)
            size_kb = out_path.stat().st_size // 1024
            ok(f"{fmt.upper()} saved → {c(str(out_path), BOLD)}  {c(f'({size_kb} KB)', DIM)}")
            generated.append(out_path)
        except Exception as ex:
            err(f"{fmt.upper()} generation failed: {ex}")
            import traceback; traceback.print_exc()

    # ── Done ──
    print()
    print(c("  ╔══════════════════════════════════════════════════════════╗", GREEN, BOLD))
    print(c("  ║   ✅  Report generation complete!                        ║", GREEN, BOLD))
    print(c("  ╚══════════════════════════════════════════════════════════╝", GREEN, BOLD))
    print()
    for p in generated:
        print(f"    📎  {c(str(p), CYAN)}")
    print()

print("henil")
if __name__ == "__main__":
    main()
    